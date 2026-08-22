"""
Reads a campaign brief .docx and returns clean text for brief_parser.py.

REWRITTEN: originally used `pandoc` (a sandbox-only tool, not something a
teammate should need to install on Windows). This version uses
`python-docx` instead -- pure Python, `pip install python-docx`, identical
behavior on Windows/Mac/Linux.

REAL BRIEFS ARE RICHER THAN OUR SYNTHETIC TEST CASES (found via
campaign_1.docx): they use RFP-style structured sections, name specific
screen-type preferences ("metro platform boards"), and -- critically --
state EXPLICIT EXCLUSIONS ("exclude bus-rear screens and value-tier
inventory in high-density residential areas"). Our original CampaignSpec
had no field for this at all (now fixed in brief_parser.py).

WHAT THIS EXTRACTS, IN TRUE DOCUMENT ORDER
python-docx's top-level `doc.paragraphs` and `doc.tables` are two SEPARATE
lists -- iterating them independently loses the original interleaving (a
table in the middle of the doc would get pulled to the end). We walk
`doc.element.body` directly instead, so a paragraph-table-paragraph
sequence in the source Word file comes out in the same order here.

Tables render as "Label: Value" lines (the brief's key facts -- Company
Name, Budget, etc. -- live in a 2-column table, and this format is what an
LLM parses most reliably). Images are skipped entirely.

HOW TO RUN
    python scoring/load_brief_docx.py data/campaign_briefs/campaign_1.docx
"""
import sys
from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph


def iter_block_items(doc):
    """Yield paragraphs and tables in the order they actually appear."""
    body = doc.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, doc)
        elif child.tag == qn("w:tbl"):
            yield Table(child, doc)


def render_table(table) -> list:
    lines = []
    for row in table.rows:
        cells = [c.text.strip() for c in row.cells]
        # Collapse repeated text from merged Word cells (python-docx returns
        # the same string for every physical cell spanned by a merge).
        seen = []
        for c in cells:
            if c and c not in seen:
                seen.append(c)
        if len(seen) >= 2:
            lines.append(f"{seen[0]}: {': '.join(seen[1:])}")
        elif seen:
            lines.append(seen[0])
    return lines


def load_brief_text(docx_path: str) -> str:
    doc = Document(docx_path)
    lines = []

    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            text = block.text.strip()
            if text:
                lines.append(text)
        elif isinstance(block, Table):
            lines.extend(render_table(block))

    return "\n".join(lines)


if __name__ == "__main__":
    path = sys.argv[1] if len(sys.argv) > 1 else "data/campaign_briefs/campaign_1.docx"
    print(load_brief_text(path))